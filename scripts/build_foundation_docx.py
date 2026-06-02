from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "foundation"
DOCX_PATH = OUT / "Correspondence_Computing_Foundations_v0_1.docx"
PDF_PATH = OUT / "Correspondence_Computing_Foundations_v0_1.pdf"
MD_PATH = OUT / "FOUNDATIONS.md"


BLUE = RGBColor(31, 77, 120)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 14, 6),
        ("Heading 2", 13, BLUE, 10, 4),
        ("Heading 3", 12, RGBColor(11, 37, 69), 8, 3),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.right_indent = Inches(0.15)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    for line in text.strip("\n").splitlines():
        run = p.add_run(line + "\n")
        run.font.name = "Consolas"
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(30, 30, 30)
    p._p.get_or_add_pPr().append(_paragraph_shading(LIGHT_GRAY))


def _paragraph_shading(fill: str):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    return shd


def add_key_value_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Inches(2.0)
    table.columns[1].width = Inches(4.25)
    hdr = table.rows[0].cells
    hdr[0].text = "Layer"
    hdr[1].text = "Role"
    for cell in hdr:
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
        for cell in cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build_docx() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Correspondence Computing Foundations")
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)
    title.paragraph_format.space_after = Pt(2)

    subtitle = doc.add_paragraph()
    subtitle.add_run("CoLean / CoPU / C-IR, v0.1").italic = True
    subtitle.paragraph_format.space_after = Pt(12)

    doc.add_heading("Executive Thesis", level=1)
    doc.add_paragraph(
        "The core contribution of this project is not a Lean agent. Lean is an "
        "important verification environment and benchmark surface, but the deeper "
        "proposal is a mathematical and computational substrate for next-generation "
        "AI reasoning."
    )
    add_code_block(
        doc,
        "Reasoning is not primarily dense tensor evaluation or one-shot text generation.\n"
        "Reasoning is structured correspondence search over composable event spaces.",
    )

    doc.add_heading("Primitive Object", level=1)
    doc.add_paragraph("The primitive object is a weighted correspondence:")
    add_code_block(doc, "X <- E -> Y")
    add_bullets(
        doc,
        [
            "X is a source space.",
            "Y is a target space.",
            "E is an event space.",
            "Each event carries source, target, weight, and evidence.",
        ],
    )
    doc.add_paragraph(
        "Instead of collapsing intermediate structure into a vector, matrix, "
        "embedding, or prompt, correspondence computing preserves the event space."
    )

    doc.add_heading("Composition And Acceleration", level=1)
    doc.add_paragraph("Given two correspondences:")
    add_code_block(doc, "X <- E -> Y\nY <- F -> Z")
    doc.add_paragraph("the composed correspondence is built by fiber join:")
    add_code_block(doc, "X <- E x_Y F -> Z")
    doc.add_paragraph(
        "The computational advantage is that only fiber-compatible events are "
        "enumerated. A dense materialization would inspect |X| * |Y| * |Z| cells; "
        "the correspondence interpreter instead evaluates sum_y |E_y| * |F_y|."
    )

    doc.add_heading("Push-Reduce And Feedback", level=1)
    doc.add_paragraph(
        "After local composition, many paths can share an endpoint pair. Push-reduce "
        "aggregates those paths while retaining evidence and weight information."
    )
    add_code_block(
        doc,
        "failed candidate -> failure class -> repair action -> updated correspondence weight",
    )
    doc.add_paragraph(
        "This makes failure useful. A failed proof, failed retrieval, failed rewrite, "
        "or failed mapping becomes a structured event rather than a discarded prompt."
    )

    doc.add_heading("Lean As Verification Surface", level=1)
    add_key_value_table(
        doc,
        [
            ("LLM", "proposal and ranking signal"),
            ("CoLean", "structured correspondence search"),
            ("Lean/Mathlib", "correctness oracle"),
        ],
    )
    doc.add_paragraph(
        "The current prototype shows that free local LLM generation fails on Lean "
        "proofs, while CoLean-guided candidate ranking recovers verified proofs. "
        "Lean is therefore a benchmark and verification case study, not the whole project."
    )

    doc.add_heading("CoPU And C-IR", level=1)
    doc.add_paragraph(
        "The long-term hardware idea is CoPU: a correspondence processing unit. "
        "A CoPU-like architecture would specialize in sparse event storage, fiber "
        "join, local composition, push-reduce, weight update, and verifier-feedback "
        "integration."
    )
    doc.add_paragraph(
        "C-IR is the proposed intermediate representation for compiling reasoning "
        "tasks into correspondence graphs and execution plans."
    )
    add_code_block(
        doc,
        "problem structure -> correspondence graph -> C-IR -> CoPU/runtime execution -> verified result",
    )

    doc.add_heading("Current Evidence", level=1)
    add_key_value_table(
        doc,
        [
            ("Structured relation benchmark", "8000x candidate-enumeration reduction"),
            ("Scale sweep", "500x -> 72000x reduction under structured sparsity"),
            ("Local LLM free generation", "0/4 Lean proofs accepted"),
            ("Local LLM + CoLean ranking", "4/4 Lean proofs accepted"),
            ("KakeyaToy.lean", "1 definition, 7 theorems, sorry = 0"),
        ],
    )

    doc.add_heading("Benchmark Path", level=1)
    add_bullets(
        doc,
        [
            "Lean accepts final files.",
            "sorry = 0.",
            "top-1 and top-k pass rates are reported.",
            "time to first verified proof is recorded.",
            "proof repair iterations are measured.",
            "Mathlib retrieval precision and dependency graph coverage are tracked.",
        ],
    )

    doc.add_heading("Summary", level=1)
    add_code_block(
        doc,
        "Do not collapse reasoning into a single vector, matrix, or prompt.\n"
        "Preserve correspondences, compose them by fibers, reduce with evidence,\n"
        "and update the graph from verifier feedback.",
    )
    doc.add_paragraph(
        "CoLean is the first public prototype. The core project is correspondence "
        "computing as a mathematical basis for future AI reasoning systems and hardware."
    )

    doc.save(DOCX_PATH)


def _pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            "FoundationTitle",
            parent=styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=24,
            leading=28,
            textColor=colors.HexColor("#0B2545"),
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            "FoundationH1",
            parent=styles["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=16,
            leading=19,
            textColor=colors.HexColor("#1F4D78"),
            spaceBefore=12,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            "FoundationBody",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=10.5,
            leading=14,
            spaceAfter=7,
        )
    )
    styles.add(
        ParagraphStyle(
            "FoundationCode",
            parent=styles["Code"],
            fontName="Courier",
            fontSize=9,
            leading=11,
            backColor=colors.HexColor("#F4F6F9"),
            borderColor=colors.HexColor("#D8DEE8"),
            borderWidth=0.5,
            borderPadding=6,
            leftIndent=8,
            rightIndent=8,
            spaceBefore=3,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            "FoundationBullet",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=10.5,
            leading=14,
            leftIndent=18,
            firstLineIndent=-10,
            spaceAfter=4,
        )
    )
    return styles


def _escape_pdf_text(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def build_pdf() -> None:
    styles = _pdf_styles()
    story = [
        Paragraph("Correspondence Computing Foundations", styles["FoundationTitle"]),
        Paragraph("CoLean / CoPU / C-IR, v0.1", styles["FoundationBody"]),
        Spacer(1, 0.08 * inch),
    ]

    lines = MD_PATH.read_text(encoding="utf-8").splitlines()
    in_code = False
    code_lines: list[str] = []
    paragraph_lines: list[str] = []

    def flush_paragraph() -> None:
        if not paragraph_lines:
            return
        text = " ".join(paragraph_lines).strip()
        if text:
            story.append(Paragraph(_escape_pdf_text(text), styles["FoundationBody"]))
        paragraph_lines.clear()

    def flush_code() -> None:
        if not code_lines:
            return
        code = "<br/>".join(_escape_pdf_text(line) for line in code_lines)
        story.append(Paragraph(code, styles["FoundationCode"]))
        code_lines.clear()

    for raw in lines:
        line = raw.rstrip()
        if line.startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                flush_paragraph()
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        if line.startswith("# "):
            continue
        if line.startswith("## "):
            flush_paragraph()
            story.append(Paragraph(_escape_pdf_text(line[3:]), styles["FoundationH1"]))
            continue
        if line.startswith("- "):
            flush_paragraph()
            story.append(Paragraph("&bull; " + _escape_pdf_text(line[2:]), styles["FoundationBullet"]))
            continue
        if not line:
            flush_paragraph()
            continue
        paragraph_lines.append(line)

    flush_paragraph()
    flush_code()

    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=LETTER,
        rightMargin=0.85 * inch,
        leftMargin=0.85 * inch,
        topMargin=0.85 * inch,
        bottomMargin=0.85 * inch,
        title="Correspondence Computing Foundations",
        author="CoLean",
    )
    doc.build(story)


def build() -> None:
    build_docx()
    build_pdf()
    print(DOCX_PATH)
    print(PDF_PATH)


if __name__ == "__main__":
    build()
